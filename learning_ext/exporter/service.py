"""导出模块：导出 Anki 牌组 / Markdown 笔记 / PDF 学习报告。"""

from __future__ import annotations

import io
import json
import zipfile
from datetime import datetime, timezone
from html import escape
from typing import Optional

from sqlalchemy.orm import Session
from sqlmodel import select

from learning_ext.db.models import (
    Card,
    KnowledgeNode,
    LearningProject,
    QuizAttempt,
    QuizQuestion,
    ReviewLog,
)
from learning_ext.progress.service import get_project_overview
from learning_ext.progress.study import sort_nodes_by_code


def export_anki_apkg(session: Session, project_id: int) -> bytes:
    """导出 Anki .apkg 牌组 (简化版：生成可导入 Anki 的 TSV + 媒体 zip)。

    Anki 原生 .apkg 是 SQLite 二进制较复杂，这里生成 Anki 可直接导入的
    "文本分隔符"格式 + 打包成 zip。用户在 Anki 里 文件→导入 即可。
    """
    cards = session.exec(select(Card).where(Card.project_id == project_id)).all()

    # Anki TSV 格式: 正面<TAB>背面<TAB>标签
    lines = ["#separator:tab\n#html:true\n"]
    for c in cards:
        tags = f"project_{project_id}"
        if c.node_id:
            tags += f" node_{c.node_id}"
        line = f"{c.front}\t{c.back}\t{tags}\n"
        lines.append(line)

    tsv_content = "".join(lines)

    # 打包成 zip
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"project_{project_id}_cards.txt", tsv_content)
        zf.writestr(
            "README.txt",
            "在 Anki 中：文件 → 导入 → 选择 txt 文件。\n"
            "分隔符选 Tab，字段映射：1=正面, 2=背面, 3=标签。\n",
        )
    return buf.getvalue()


def export_markdown(session: Session, project_id: int) -> str:
    """导出完整学习笔记 Markdown。"""
    project = session.get(LearningProject, project_id)
    if project is None:
        raise ValueError(f"Project {project_id} not found")

    nodes = session.exec(
        select(KnowledgeNode).where(KnowledgeNode.project_id == project_id)
    ).all()
    nodes = sort_nodes_by_code(list(nodes))
    cards = session.exec(select(Card).where(Card.project_id == project_id)).all()

    md = f"# {project.title}\n\n"
    md += f"> **选题**: {project.topic}\n\n"
    md += f"> **目标**: {project.goal}\n\n"
    md += f"> **导出时间**: {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')}\n\n"
    md += "---\n\n"

    # 知识图谱
    md += "## 知识图谱\n\n"
    overview = get_project_overview(session, project_id)
    md += (
        f"- 知识点总数: {overview['total']}\n"
        f"- 已掌握: {overview['mastered']} ({overview['mastered_pct']:.0%})\n"
        f"- 平均掌握度: {overview['avg_mastery']:.0%}\n"
        f"- 预计总时长: {overview['total_hours']:.1f} 小时\n\n"
    )

    # 各知识点
    md += "## 知识点详情\n\n"
    for n in nodes:
        status_emoji = {
            "mastered": "✅",
            "reviewing": "🔄",
            "learning": "📖",
            "weak": "⚠️",
            "pending": "⏳",
        }.get(n.status, "⏳")
        md += f"### {status_emoji} [{n.code}] {n.title}\n\n"
        md += f"- 阶段: {n.stage} | 难度: {n.difficulty}/5 | 掌握度: {n.mastery:.0%}\n"
        md += f"- 预计学时: {n.est_hours}h\n\n"
        if n.description:
            md += f"{n.description}\n\n"

        # 该节点的卡片
        node_cards = [c for c in cards if c.node_id == n.id]
        if node_cards:
            md += "**复习卡片**:\n\n"
            for c in node_cards:
                md += f"- **Q**: {c.front}\n  **A**: {c.back}\n\n"
        md += "---\n\n"

    return md


def export_learning_plan_docx(session: Session, project_id: int) -> bytes:
    """导出 Word 学习计划文档 (.docx)."""
    project = session.get(LearningProject, project_id)
    if project is None:
        raise ValueError(f"Project {project_id} not found")

    nodes = session.exec(
        select(KnowledgeNode).where(KnowledgeNode.project_id == project_id)
    ).all()
    nodes = sort_nodes_by_code(list(nodes))
    stages = _load_project_stages(project, nodes)
    total_hours = sum(float(n.est_hours or 0) for n in nodes)

    try:
        return _export_learning_plan_docx_with_python_docx(
            project, nodes, stages, total_hours
        )
    except ModuleNotFoundError:
        return _export_learning_plan_docx_minimal(project, nodes, stages, total_hours)


def _load_project_stages(
    project: LearningProject, nodes: list[KnowledgeNode]
) -> list[dict]:
    try:
        roadmap = json.loads(project.roadmap_json or "{}")
    except json.JSONDecodeError:
        roadmap = {}
    stages = roadmap.get("stages") if isinstance(roadmap, dict) else None
    if isinstance(stages, list) and stages:
        return [s for s in stages if isinstance(s, dict)]

    seen = []
    for node in nodes:
        if node.stage not in [s["stage"] for s in seen]:
            seen.append({"name": node.stage, "stage": node.stage, "goal": ""})
    return seen


def _export_learning_plan_docx_with_python_docx(
    project: LearningProject,
    nodes: list[KnowledgeNode],
    stages: list[dict],
    total_hours: float,
) -> bytes:
    from docx import Document

    doc = Document()
    doc.add_heading(project.title, level=0)
    doc.add_paragraph("学习计划文档")

    doc.add_heading("学习目标", level=1)
    doc.add_paragraph(project.goal or "掌握该主题的核心知识并能应用。")
    doc.add_heading("学习者背景", level=1)
    doc.add_paragraph(project.background or "未填写")

    doc.add_heading("路线总览", level=1)
    doc.add_paragraph(
        f"总计 {len(nodes)} 个知识点，预计 {total_hours:.0f} 小时；"
        f"建议每周投入 {project.weekly_hours:.0f} 小时。"
    )

    doc.add_heading("三阶段安排", level=1)
    stage_table = doc.add_table(rows=1, cols=4)
    stage_table.style = "Table Grid"
    for i, header in enumerate(["阶段", "目标", "节点数", "预计小时"]):
        stage_table.rows[0].cells[i].text = header
    for stage in stages:
        code = stage.get("stage", "")
        row = stage_table.add_row().cells
        row[0].text = str(stage.get("name") or code)
        row[1].text = str(stage.get("goal") or "")
        row[2].text = str(sum(1 for n in nodes if n.stage == code))
        row[3].text = f"{sum(n.est_hours for n in nodes if n.stage == code):.0f}h"

    doc.add_heading("学习路线", level=1)
    node_table = doc.add_table(rows=1, cols=6)
    node_table.style = "Table Grid"
    for i, header in enumerate(["编号", "阶段", "标题", "学时", "难度", "状态"]):
        node_table.rows[0].cells[i].text = header
    for node in nodes:
        row = node_table.add_row().cells
        row[0].text = node.code
        row[1].text = node.stage
        row[2].text = node.title
        row[3].text = f"{node.est_hours:.0f}h"
        row[4].text = str(node.difficulty)
        row[5].text = node.status

    doc.add_heading("最终交付物", level=1)
    for item in _learning_plan_deliverables():
        doc.add_paragraph(item, style="List Bullet")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _export_learning_plan_docx_minimal(
    project: LearningProject,
    nodes: list[KnowledgeNode],
    stages: list[dict],
    total_hours: float,
) -> bytes:
    body = [
        _docx_paragraph(project.title, "Title"),
        _docx_paragraph("学习计划文档", "Subtitle"),
        _docx_paragraph("学习目标", "Heading1"),
        _docx_paragraph(project.goal or "掌握该主题的核心知识并能应用。"),
        _docx_paragraph("学习者背景", "Heading1"),
        _docx_paragraph(project.background or "未填写"),
        _docx_paragraph("路线总览", "Heading1"),
        _docx_paragraph(
            f"总计 {len(nodes)} 个知识点，预计 {total_hours:.0f} 小时；"
            f"建议每周投入 {project.weekly_hours:.0f} 小时。"
        ),
        _docx_paragraph("三阶段安排", "Heading1"),
        _docx_table(
            [["阶段", "目标", "节点数", "预计小时"]]
            + [
                [
                    str(stage.get("name") or stage.get("stage") or ""),
                    str(stage.get("goal") or ""),
                    str(sum(1 for n in nodes if n.stage == stage.get("stage"))),
                    f"{sum(n.est_hours for n in nodes if n.stage == stage.get('stage')):.0f}h",
                ]
                for stage in stages
            ]
        ),
        _docx_paragraph("学习路线", "Heading1"),
        _docx_table(
            [["编号", "阶段", "标题", "学时", "难度", "状态"]]
            + [
                [
                    n.code,
                    n.stage,
                    n.title,
                    f"{n.est_hours:.0f}h",
                    str(n.difficulty),
                    n.status,
                ]
                for n in nodes
            ]
        ),
        _docx_paragraph("最终交付物", "Heading1"),
    ]
    body.extend(_docx_paragraph(item) for item in _learning_plan_deliverables())
    return _build_minimal_docx("".join(body))


def _learning_plan_deliverables() -> list[str]:
    return [
        "一个能接收技术问题、检索资料、记录证据并输出引用答案的 research agent。",
        "一个可复用的证据记录 skill。",
        "一个安全 MCP server 接入示例，外部动作前需要批准。",
        "一个 critic agent 审核流程，覆盖证据完整性、风险门控和工具失败回退。",
        "一份故障模式记录，说明超时、格式错误、服务不可用时的检测与降级策略。",
    ]


def _docx_paragraph(text: str = "", style: Optional[str] = None) -> str:
    style_xml = f'<w:pPr><w:pStyle w:val="{style}"/></w:pPr>' if style else ""
    return (
        f"<w:p>{style_xml}<w:r><w:t xml:space=\"preserve\">"
        f"{escape(str(text))}</w:t></w:r></w:p>"
    )


def _docx_table(rows: list[list[str]]) -> str:
    parts = [
        "<w:tbl><w:tblPr><w:tblW w:w=\"0\" w:type=\"auto\"/>"
        "<w:tblBorders>"
        "<w:top w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"BFC7D5\"/>"
        "<w:left w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"BFC7D5\"/>"
        "<w:bottom w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"BFC7D5\"/>"
        "<w:right w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"BFC7D5\"/>"
        "<w:insideH w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"BFC7D5\"/>"
        "<w:insideV w:val=\"single\" w:sz=\"4\" w:space=\"0\" w:color=\"BFC7D5\"/>"
        "</w:tblBorders></w:tblPr>"
    ]
    for row in rows:
        parts.append("<w:tr>")
        for cell in row:
            parts.append(
                '<w:tc><w:tcPr><w:tcW w:w="2400" w:type="dxa"/></w:tcPr>'
            )
            parts.append(_docx_paragraph(cell))
            parts.append("</w:tc>")
        parts.append("</w:tr>")
    parts.append("</w:tbl>")
    return "".join(parts)


def _build_minimal_docx(body_xml: str) -> bytes:
    content_types = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/word/document.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.document.main+xml"/><Override PartName="/word/styles.xml" ContentType="application/vnd.openxmlformats-officedocument.wordprocessingml.styles+xml"/></Types>"""
    rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="word/document.xml"/></Relationships>"""
    doc_rels = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/styles" Target="styles.xml"/></Relationships>"""
    styles = """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:styles xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:style w:type="paragraph" w:default="1" w:styleId="Normal"><w:name w:val="Normal"/><w:rPr><w:rFonts w:ascii="Aptos" w:eastAsia="Microsoft YaHei"/><w:sz w:val="22"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Title"><w:name w:val="Title"/><w:rPr><w:b/><w:sz w:val="40"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Subtitle"><w:name w:val="Subtitle"/><w:rPr><w:color w:val="64748B"/><w:sz w:val="26"/></w:rPr></w:style><w:style w:type="paragraph" w:styleId="Heading1"><w:name w:val="heading 1"/><w:rPr><w:b/><w:sz w:val="30"/><w:color w:val="1F4E79"/></w:rPr></w:style></w:styles>"""
    document = (
        """<?xml version="1.0" encoding="UTF-8" standalone="yes"?>
<w:document xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><w:body>"""
        + body_xml
        + '<w:sectPr><w:pgSz w:w="11906" w:h="16838"/><w:pgMar w:top="1440" w:right="1440" w:bottom="1440" w:left="1440"/></w:sectPr></w:body></w:document>'
    )

    buf = io.BytesIO()
    with zipfile.ZipFile(buf, "w", zipfile.ZIP_DEFLATED) as zf:
        zf.writestr("[Content_Types].xml", content_types)
        zf.writestr("_rels/.rels", rels)
        zf.writestr("word/_rels/document.xml.rels", doc_rels)
        zf.writestr("word/document.xml", document)
        zf.writestr("word/styles.xml", styles)
    return buf.getvalue()


def export_progress_report(session: Session, project_id: int) -> str:
    """导出 PDF 友好的学习进度报告 (HTML 格式，可浏览器打印为 PDF)。"""
    project = session.get(LearningProject, project_id)
    if project is None:
        raise ValueError(f"Project {project_id} not found")

    nodes = session.exec(
        select(KnowledgeNode).where(KnowledgeNode.project_id == project_id)
    ).all()
    nodes = sort_nodes_by_code(list(nodes))
    overview = get_project_overview(session, project_id)

    html = f"""<!DOCTYPE html>
<html lang="zh-CN"><head><meta charset="utf-8">
<title>学习报告 - {project.title}</title>
<style>
body {{ font-family: -apple-system, "Microsoft YaHei", sans-serif; max-width: 800px; margin: 40px auto; padding: 20px; color: #333; }}
h1 {{ border-bottom: 3px solid #4F46E5; padding-bottom: 10px; }}
.metric {{ display: inline-block; background: #F3F4F6; padding: 12px 20px; margin: 5px; border-radius: 8px; }}
.metric .val {{ font-size: 24px; font-weight: bold; color: #4F46E5; }}
.node {{ margin: 12px 0; padding: 12px; border-left: 4px solid #ddd; background: #fafafa; }}
.node.mastered {{ border-color: #10B981; }}
.node.weak {{ border-color: #EF4444; }}
.bar {{ height: 8px; background: #E5E7EB; border-radius: 4px; margin-top: 4px; }}
.bar > div {{ height: 100%; background: #4F46E5; border-radius: 4px; }}
</style></head><body>
<h1>📊 学习进度报告</h1>
<p><strong>{project.title}</strong></p>
<p>选题：{project.topic}</p>

<div>
  <div class="metric"><div class="val">{overview["total"]}</div>知识点</div>
  <div class="metric"><div class="val">{overview["mastered"]}</div>已掌握</div>
  <div class="metric"><div class="val">{overview["mastered_pct"]:.0%}</div>完成率</div>
  <div class="metric"><div class="val">{overview["avg_mastery"]:.0%}</div>平均掌握</div>
</div>

<h2>知识点掌握详情</h2>
"""
    for n in nodes:
        html += f"""<div class="node {n.status}">
  <strong>[{n.code}] {n.title}</strong>
  <span style="float:right;color:#888">{n.status}</span>
  <div class="bar"><div style="width:{n.mastery * 100:.0f}%"></div></div>
  <small>掌握度 {n.mastery:.0%} | 难度 {n.difficulty}/5 | {n.est_hours}h</small>
</div>
"""
    html += f"\n<p style='text-align:center;color:#999;margin-top:40px;'>生成于 {datetime.now(timezone.utc).strftime('%Y-%m-%d %H:%M')} | 学习 Agent</p>\n"
    html += "</body></html>"
    return html
